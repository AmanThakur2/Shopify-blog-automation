import json
import datetime
import docx
import json
from docx import Document
from google.colab import files

##########################################################
#
## Generate HTML File using JSON file
#
def generate_hightlight(blogs_map):
    generate_hightlight_str = ''
    highlight_list = [highlights_anchor.replace("[[HIGHLIGHT_SECTION]]","section"+str(idx+1)).replace("[[HIGHLIGHT_NAME]]",highlights) for idx,highlights in enumerate(blogs_map)]
    ordered_highlight_list = "<ol>" + "".join(highlight_list)+"</ol>"
    generate_hightlight_str =  generate_hightlight_str + highlights_header + ordered_highlight_list
    generate_hightlight_str = div_highlight_container.replace("[[KEY_HIGHLIGHT]]",generate_hightlight_str)
    return generate_hightlight_str


def generate_sections(blogs_map):
    generate_section_str = []
    index = 0

    for sections in blogs_map:

        index += 1
        sect_header = section_header.replace("[[SECTION_NAME]]",sections["sec_title"]).replace("[[SECTION_NO]]",str(index))
        # sect_img = "".join([section_img.replace("[[SECTION_IMAGE]]",img_url) for img_url in sections.get("sec_img") ]) if sections.get("sec_img") else ""

        final = ''

        for sect_details in sections["sec_content"]:
            para_str ,point_para_str,inner_list ,para_content_str,img_para_str, outer_lst='','','' , '', '',''

            if sect_details["type"] == "paragraph":
                para_str =  "<p>"+sect_details["content"]+ "</p>"
                if sect_details.get("para_content"):
                    for points_list in  sect_details.get("para_content"):
                        if points_list.get("list") :
                            outer_lst  = "<li>"+points_list.get("list")+"</li>"
                        if points_list.get("sub_list") :
                            inner_list = "<ul><li>" + "</li><li>".join(points_list.get("sub_list"))+ "</li></ul>"
                        para_content_str =para_content_str + section_point_list_tag.replace("[[UL_TAG_PARA]]",outer_lst + inner_list)
                para_str = para_str + para_content_str

            if sect_details["type"] == "point" :
                point_para_str = '<li>' +sect_details["content"] +'</li>'
                if sect_details.get("para_content"):
                    for points_list in  sect_details.get("para_content"):
                        if points_list.get("list") :
                            outer_lst  = "<li>"+points_list.get("list")+"</li>"
                        if points_list.get("sub_list") :
                            inner_list = "<ul><li>" + "</li><li>".join(points_list.get("sub_list"))+ "</li></ul>"
                        para_content_str =para_content_str + '<ul>'+outer_lst + inner_list+'</ul>'
                point_para_str = section_point_para_str.replace("[[COMBINE_NAME]]",point_para_str + para_content_str)

            if sect_details["type"] == "img" :
                img_para_str = section_img.replace("[[SECTION_IMAGE]]",sect_details["content"] )

            final = final + para_str + point_para_str + img_para_str
        generate_str =  sect_header + final
        generate_section_str.append(generate_str)

    return generate_section_str


def Read_doc(file_name,data):
    doc = docx.Document(file_name)
    new_sec = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("###KEY_HIGHLIGHT_START###"):
            if "key_highlights" not in data:
                data["key_highlights"] = []
            data["key_highlights"].append("".join(text.split("###KEY_HIGHLIGHT_START###")[1:]).strip())

        elif text.startswith("###SECTION_TITLE###"):
            if new_sec:
                data["sections"].append(new_sec)
            new_sec = {"sec_title": "".join(text.split("###SECTION_TITLE###")[1:]).strip(),"sec_content": []}

        elif text.startswith("###PARAGRAPH###"):
            if new_sec:
                paragraph = {
                    "type": "paragraph",
                    "content": "".join(text.split("###PARAGRAPH###")).strip(),
                    "para_content": []
                }
                new_sec["sec_content"].append(paragraph)

        elif text.startswith("###SEC_IMAGE###"):
            if new_sec:
                paragraph = {
                    "type" : "img",
                    "content" : "".join(text.split("###SEC_IMAGE###")).strip()
                }
                new_sec["sec_content"].append(paragraph)

        elif text.startswith("###POINT###"):
            if new_sec:

                points = {
                    "type" : "point",
                    "content" : "".join(text.split("###POINT###")).strip(),"para_content" :[]
                }
                new_sec["sec_content"].append(points)

        elif text.startswith("###SUB_POINT###"):
            if new_sec:
                lst = {
                    "list" : "".join(text.split("###SUB_POINT###")).strip()
                }
                new_sec["sec_content"][0]["para_content"].append(lst)
    if new_sec:
        data["sections"].append(new_sec)

    return data

##########################################################
#
## STATIC VARIABLES
#
default_css = '''
<style>
    <!--
    .blog-img {
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
        margin-top:20px;
        margin-bottom:20px;

    }

    .blog-img span {
        text-align: center;
    }

    html {
        scroll-behavior: smooth;
    }

    ol li {
        list-style: none;
        position: relative;
    }

    ol li::before {
        content: '';
        left: -13px;
        top: 13px;
        position: absolute;
        height: 5px;
        width: 5px;
        border-radius: 50%;
        background-color: #9F4A17;
    }

    .blog-content ol li::before {
        top: 19px;
    }

    .highlight-container {
        border: 1px solid #9F4A17;
        padding: 10px 20px;
        border-radius: 10px;
        margin-bottom: 20px;
    }

    .highlight-container h1 {
        margin: 10px 0px;
    }

    p {
        text-align: justify;
    }

    ul li {
        list-style: none;
        text-align: justify;
    }

    ul li .list-content {
        padding-left: 15px;
        text-align: justify;
    }

    .numerical-list {
        list-style: auto !important;
    }

    h3 {
        font-size: 25px !important;
    }

    .product-list li::before {
        top: 41px;
    }
    -->
</style>
'''

### DIV class variable
# static variable
div_blog_container = '''<div class="blog-container">[[INNER_ELEMENT]]</div>'''
div_highlight_container  = '''<div class="highlight-container">[[KEY_HIGHLIGHT]]</div>'''
div_blog_img = '''<div class="blog-img">'''
div_point = '''<div class="point">'''
# dynamic variable
div_section_id =  '''<div class="s[[SECTION_ID]]">'''


### HIGHLIGHT
highlights_header = '''<h1><b>Highlights</b></h1>'''
highlights_anchor= '''<li><a href="#[[HIGHLIGHT_SECTION]]">[[HIGHLIGHT_NAME]]</a></li>'''

### SECTIONS
section_div_class = '''<div class="s[[SECTION_ID]]">[[SECTION_DETAIL]]</div>'''
section_header = '''<h2 style="margin-top: 25px; margin-bottom: 10px" id="section[[SECTION_NO]]"><b>[[SECTION_NAME]]</b></h2>'''
section_points_para = '''<div class="point"><ul>[[SECTION_POINT_PARA]]</ul></div>'''
section_img = '''<br><div class="blog-img"><img src="[[SECTION_IMAGE]]"/></div><br>'''
## Point List Section
section_point_para_str = '<ul style="list-style-type: disc">[[COMBINE_NAME]]</ul>'
section_point_list_tag = '''<ul style="list-style-type: disc">[[UL_TAG_PARA]]</ul>'''

##########################################################
#
## Final Run
#
def main():

    upload_doc  = files.upload()
    doc_file = 'blog_test.docx'

    doc = Document(doc_file)

    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold and run.italic:
                run.text = f"<b><i>{run.text}</i></b>"
            elif run.italic:

                run.text = f"<i>{run.text}</i>"
            elif run.bold:

                run.text = f"<b>{run.text}</b>"
    doc.save('blog_test.docx')


    data = {
        "key_highlights" : [],
        "sections" : [],
    }

    print(json.dumps(Read_doc(doc_file,data),indent=4))

    #data = Read_doc(doc_file,data)
    with open('data.json', 'w') as f:
        json.dump(data, f)


    json_file = 'data.json'

    with open(json_file, 'r') as file:
        blogs_map = json.load(file)
    # blogs_map  = json.loads(json_file)

    # Generate Key Highlights
    hightlight_str = generate_hightlight(blogs_map["key_highlights"])
    # Generate Sections
    sections_str = "".join(generate_sections(blogs_map["sections"]))
    merge_all_elmt= default_css+ hightlight_str + sections_str


    final_html= div_blog_container.replace("[[INNER_ELEMENT]]",merge_all_elmt)

    with open('index2.html', 'w') as file:
        file.write(final_html)